"""
Generate the ChatFlow360 Knowledge Base Onboarding Questionnaire (EN/ES)
Output: docs/ChatFlow360-Knowledge-Questionnaire.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Style overrides --
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

TEAL = RGBColor(0x2F, 0x92, 0xAD)
DARK_NAVY = RGBColor(0x0F, 0x1C, 0x2E)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_BG = RGBColor(0xF0, 0xF4, 0xF7)


def add_colored_heading(text, level=1, color=DARK_NAVY):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
    return heading


def add_section_number(number, title_en, title_es):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)

    run_num = p.add_run(f"{number}. ")
    run_num.font.size = Pt(14)
    run_num.font.bold = True
    run_num.font.color.rgb = TEAL

    run_en = p.add_run(f"{title_en}")
    run_en.font.size = Pt(14)
    run_en.font.bold = True
    run_en.font.color.rgb = DARK_NAVY

    run_sep = p.add_run("  /  ")
    run_sep.font.size = Pt(12)
    run_sep.font.color.rgb = GRAY

    run_es = p.add_run(f"{title_es}")
    run_es.font.size = Pt(12)
    run_es.font.italic = True
    run_es.font.color.rgb = GRAY


def add_question(en, es):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    run_en = p.add_run(f"  {en}")
    run_en.font.size = Pt(10.5)
    run_en.font.bold = True
    run_en.font.color.rgb = DARK_NAVY

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)

    run_es = p2.add_run(f"  {es}")
    run_es.font.size = Pt(10)
    run_es.font.italic = True
    run_es.font.color.rgb = GRAY


def add_answer_box(lines=4):
    """Add a light-bordered answer area"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    # Set cell height via paragraph spacing
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(lines * 14)

    # Light border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.makeelement(qn("w:tcBorders"), {})
    for border_name in ["top", "left", "bottom", "right"]:
        border = borders.makeelement(
            qn(f"w:{border_name}"),
            {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "0",
                qn("w:color"): "B0BFCB",
            },
        )
        borders.append(border)
    tcPr.append(borders)

    # Light background
    shading = tcPr.makeelement(
        qn("w:shd"), {qn("w:fill"): "F8FAFB", qn("w:val"): "clear"}
    )
    tcPr.append(shading)

    doc.add_paragraph()  # spacer


def add_tip(text_en, text_es):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)

    run_icon = p.add_run("TIP: ")
    run_icon.font.size = Pt(9)
    run_icon.font.bold = True
    run_icon.font.color.rgb = TEAL

    run_en = p.add_run(text_en)
    run_en.font.size = Pt(9)
    run_en.font.color.rgb = GRAY

    run_sep = p.add_run("  |  ")
    run_sep.font.size = Pt(9)
    run_sep.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    run_es = p.add_run(text_es)
    run_es.font.size = Pt(9)
    run_es.font.italic = True
    run_es.font.color.rgb = GRAY


# ============================================================
# COVER
# ============================================================

# Add logo if exists
logo_path = os.path.join(os.path.dirname(__file__), "..", "public", "logo.png")
if os.path.exists(logo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(2))

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run("Knowledge Base\nOnboarding Questionnaire")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = DARK_NAVY

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
run = p.add_run("Cuestionario de Base de Conocimiento")
run.font.size = Pt(16)
run.font.italic = True
run.font.color.rgb = GRAY

# Description
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
run = p.add_run(
    "Complete this questionnaire so your AI assistant can answer\n"
    "visitor questions from day one.\n\n"
    "Complete este cuestionario para que su asistente de IA pueda\n"
    "responder las preguntas de los visitantes desde el primer dia."
)
run.font.size = Pt(10.5)
run.font.color.rgb = GRAY

# Org info fields
doc.add_paragraph()
fields = [
    ("Organization / Organizacion:", ""),
    ("Contact Name / Nombre de Contacto:", ""),
    ("Date / Fecha:", ""),
]
for label, _ in fields:
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = DARK_NAVY
    run2 = p.add_run("  _______________________________________________")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

doc.add_page_break()

# ============================================================
# INSTRUCTIONS
# ============================================================

add_colored_heading("How to Use This Document", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Answer each section in the language your customers use most. "
    "If your audience is bilingual, include answers in both English and Spanish. "
    "Write naturally — your AI assistant will use this information exactly as you provide it."
)
run.font.size = Pt(10)
run.font.color.rgb = GRAY

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run = p.add_run(
    "Responda cada seccion en el idioma que mas usan sus clientes. "
    "Si su audiencia es bilingue, incluya respuestas en ingles y espanol. "
    "Escriba de forma natural — su asistente de IA usara esta informacion exactamente como la proporcione."
)
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = GRAY

doc.add_paragraph()

# ============================================================
# SECTION 1: About Your Business
# ============================================================

add_section_number(1, "About Your Business", "Sobre Su Negocio")

add_question(
    "What does your business do? Describe it in 2-3 sentences as you would tell a new customer.",
    "Que hace su negocio? Describalo en 2-3 oraciones como se lo diria a un nuevo cliente.",
)
add_answer_box(5)

add_question(
    "How many years of experience do you have?",
    "Cuantos anos de experiencia tienen?",
)
add_answer_box(2)

add_question(
    "What areas or regions do you serve?",
    "Que areas o regiones atienden?",
)
add_answer_box(2)

add_question(
    "What languages does your team speak?",
    "Que idiomas habla su equipo?",
)
add_answer_box(2)

# ============================================================
# SECTION 2: Services / Products
# ============================================================

add_section_number(2, "Services & Products", "Servicios y Productos")

add_question(
    "List your main services or products with a brief description of each.",
    "Liste sus servicios o productos principales con una breve descripcion de cada uno.",
)
add_tip(
    "Example: 'Immigration Law — We help with visa applications, green cards, and citizenship.'",
    "Ejemplo: 'Derecho Migratorio — Ayudamos con solicitudes de visa, green cards y ciudadania.'",
)
add_answer_box(8)

add_question(
    "What is your most requested service/product?",
    "Cual es su servicio/producto mas solicitado?",
)
add_answer_box(3)

add_question(
    "Are there services you DO NOT offer that people commonly ask about?",
    "Hay servicios que NO ofrecen pero que la gente pregunta frecuentemente?",
)
add_tip(
    "This helps the AI avoid making promises you can't keep.",
    "Esto ayuda a la IA a no hacer promesas que no pueden cumplir.",
)
add_answer_box(3)

# ============================================================
# SECTION 3: Pricing & Payment
# ============================================================

doc.add_page_break()
add_section_number(3, "Pricing & Payment", "Precios y Formas de Pago")

add_question(
    "What is your pricing structure? (fixed rates, hourly, free consultation, quote-based, etc.)",
    "Cual es su estructura de precios? (tarifas fijas, por hora, consulta gratis, cotizacion, etc.)",
)
add_answer_box(4)

add_question(
    "What payment methods do you accept?",
    "Que metodos de pago aceptan?",
)
add_answer_box(2)

add_question(
    "Do you offer financing or payment plans?",
    "Ofrecen financiamiento o planes de pago?",
)
add_answer_box(2)

# ============================================================
# SECTION 4: How to Get Started
# ============================================================

add_section_number(4, "How to Get Started", "Como Empezar")

add_question(
    "What is the first step for a new customer? (call, book online, visit, fill a form, etc.)",
    "Cual es el primer paso para un nuevo cliente? (llamar, agendar online, visitar, llenar formulario, etc.)",
)
add_answer_box(3)

add_question(
    "What should the customer bring or prepare for the first visit/meeting?",
    "Que debe traer o preparar el cliente para la primera visita/reunion?",
)
add_answer_box(3)

add_question(
    "How long does the typical process take from start to finish?",
    "Cuanto tiempo toma el proceso tipico de inicio a fin?",
)
add_answer_box(2)

# ============================================================
# SECTION 5: Location & Hours
# ============================================================

add_section_number(5, "Location & Hours", "Ubicacion y Horarios")

add_question(
    "What is your physical address?",
    "Cual es su direccion fisica?",
)
add_answer_box(2)

add_question(
    "What are your business hours? (include weekends if applicable)",
    "Cual es su horario de atencion? (incluya fines de semana si aplica)",
)
add_answer_box(3)

add_question(
    "Do you offer virtual/remote appointments?",
    "Ofrecen citas virtuales/remotas?",
)
add_answer_box(2)

# ============================================================
# SECTION 6: Contact Information
# ============================================================

doc.add_page_break()
add_section_number(6, "Contact Information", "Informacion de Contacto")

add_question(
    "Phone number:",
    "Numero de telefono:",
)
add_answer_box(1)

add_question(
    "Email address:",
    "Correo electronico:",
)
add_answer_box(1)

add_question(
    "Website URL:",
    "Sitio web:",
)
add_answer_box(1)

add_question(
    "Social media profiles (Instagram, Facebook, LinkedIn, etc.):",
    "Perfiles de redes sociales (Instagram, Facebook, LinkedIn, etc.):",
)
add_answer_box(2)

add_question(
    "How can customers book an appointment? (link, phone, form)",
    "Como pueden los clientes agendar una cita? (enlace, telefono, formulario)",
)
add_answer_box(2)

# ============================================================
# SECTION 7: Frequently Asked Questions
# ============================================================

add_section_number(7, "Frequently Asked Questions", "Preguntas Frecuentes")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run(
    "List the 5-10 questions your customers ask most often, with the answers you typically give."
)
run.font.size = Pt(10)
run.font.color.rgb = GRAY
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(12)
run2 = p2.add_run(
    "Liste las 5-10 preguntas que sus clientes hacen con mas frecuencia, con las respuestas que tipicamente da."
)
run2.font.size = Pt(10)
run2.font.italic = True
run2.font.color.rgb = GRAY

for i in range(1, 8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(f"Q{i} / P{i}:")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = TEAL
    add_answer_box(3)

# ============================================================
# FOOTER NOTE
# ============================================================

doc.add_page_break()
add_colored_heading("What Happens Next?", level=2)

p = doc.add_paragraph()
run = p.add_run(
    "Once you complete this questionnaire, our team will upload this information "
    "to your AI assistant's Knowledge Base. Your assistant will immediately be able "
    "to answer visitor questions based on the information you provided.\n\n"
    "You can always add, edit, or remove knowledge items later from your "
    "ChatFlow360 dashboard under Settings > AI Settings > Knowledge Base."
)
run.font.size = Pt(10.5)
run.font.color.rgb = GRAY

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run(
    "Una vez que complete este cuestionario, nuestro equipo subira esta informacion "
    "a la Base de Conocimiento de su asistente de IA. Su asistente podra responder "
    "inmediatamente las preguntas de los visitantes basandose en la informacion proporcionada.\n\n"
    "Siempre podra agregar, editar o eliminar elementos de conocimiento desde su "
    "panel de ChatFlow360 en Configuracion > Ajustes de IA > Base de Conocimiento."
)
run.font.size = Pt(10.5)
run.font.italic = True
run.font.color.rgb = GRAY

# Branding footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("chatflow360.com")
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = TEAL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI-Powered Live Chat for Miami Businesses")
run.font.size = Pt(9)
run.font.color.rgb = GRAY


# ============================================================
# SAVE
# ============================================================

output_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
output_path = os.path.join(output_dir, "ChatFlow360-Knowledge-Questionnaire.docx")
doc.save(output_path)
print(f"Document saved: {output_path}")
